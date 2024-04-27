from flask import Flask, render_template, request, send_file
from docx import Document
from io import BytesIO
import xlwt
import os
import re
import fitz
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
import nltk

# Download NLTK data
nltk.download('punkt')

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'doc', 'docx', 'pdf'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_doc(doc_file):
    doc = Document(doc_file)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def extract_data_from_cv(cv_file):
    file_extension = cv_file.filename.rsplit('.', 1)[1].lower()
    text = ""

    try:
        if file_extension == 'docx':
            text = extract_text_from_docx(cv_file)
        elif file_extension == 'doc':
            # Extract text from .doc file
            text = extract_text_from_doc(cv_file)
        elif file_extension == 'pdf':
            text = ""
            with fitz.open(stream=cv_file.read(), filetype="pdf") as pdf_file:
                for page in pdf_file:
                    text += page.get_text()
    except Exception as e:
        print(f"Error processing {cv_file.filename}: {str(e)}")
        return None

    # Generate a summary of the CV using sumy
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = LsaSummarizer()
    summary = summarizer(parser.document, sentences_count=2)

    # Extract text from the tuple and replace newline characters
    summary_text = " ".join([sentence.__str__() for sentence in summary])

    email = re.findall(r'[\w\.-]+@[\w\.-]+', text)
    phone_numbers = re.findall(r'(\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}|\(\d{3}\)\s*\d{3}[-\.\s]??\d{4}|\d{3}[-\.\s]??\d{4})', text)

    return {
        'text': summary_text,
        'email': email,
        'phone_numbers': phone_numbers
    }

def save_to_excel(data):
    workbook = xlwt.Workbook()
    sheet = workbook.add_sheet('CV Data')

    headers = ['Text', 'Email', 'Phone Numbers']
    for col, header in enumerate(headers):
        sheet.write(0, col, header)

    for row, cv_data in enumerate(data, start=1):
        sheet.write(row, 0, cv_data['text'])
        sheet.write(row, 1, ", ".join(cv_data['email']))
        sheet.write(row, 2, ", ".join(cv_data['phone_numbers']))

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'extracted_data.xls')
    workbook.save(file_path)

    return file_path

@app.route('/')
def upload_form():
    return render_template('upload.html')

@app.route('/', methods=['POST'])
def upload_cv():
    uploaded_files = request.files.getlist('file[]')
    data = []
    for file in uploaded_files:
        if file and allowed_file(file.filename):
            cv_data = extract_data_from_cv(file)
            if cv_data:
                data.append(cv_data)

    if data:
        excel_file = save_to_excel(data)
        return send_file(excel_file, as_attachment=True)

    return 'No valid CVs uploaded.'

if __name__ == "__main__":
    app.run(debug=True)
